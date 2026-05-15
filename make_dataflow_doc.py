"""Generate DevPost Agent data flow story as a Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x37, 0x5E)   # headings
TEAL   = RGBColor(0x00, 0x7A, 0x87)   # sub-headings / arrows
GRAY   = RGBColor(0x44, 0x44, 0x44)   # body
CODE_BG = RGBColor(0xF5, 0xF5, 0xF5)
CODE_FG = RGBColor(0x2D, 0x2D, 0x2D)
GREEN  = RGBColor(0x1E, 0x6E, 0x3A)
ORANGE = RGBColor(0xC2, 0x5A, 0x00)
RED    = RGBColor(0x9B, 0x1C, 0x1C)

# ── Helpers ───────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = NAVY
    p.runs[0].font.size = Pt(20)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = TEAL
    p.runs[0].font.size = Pt(14)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = NAVY
    p.runs[0].font.size = Pt(12)
    return p

def body(text, bold=False, italic=False, colour=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = colour or GRAY
    run.bold  = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    return p

def code_block(lines):
    """Render a shaded code block (each line = one paragraph)."""
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = CODE_FG
        # shade the paragraph
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F0F0F0")
        pPr.append(shd)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.4)

def arrow(text):
    p = doc.add_paragraph()
    run = p.add_run(f"  ➜  {text}")
    run.font.size = Pt(11)
    run.font.color.rgb = TEAL
    run.bold = True
    p.paragraph_format.space_after = Pt(2)
    return p

def divider():
    doc.add_paragraph("─" * 72)

def spacer():
    doc.add_paragraph("")

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# TITLE PAGE
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_para.add_run("DevPost Agent")
r.font.size  = Pt(32)
r.font.bold  = True
r.font.color.rgb = NAVY

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_para.add_run("Data Flow Story — How Your Code Travels From Git to the World")
r2.font.size   = Pt(14)
r2.font.italic = True
r2.font.color.rgb = TEAL

spacer()
divider()
spacer()

body(
    "This document is a plain-English walkthrough of exactly what happens "
    "inside DevPost Agent every single time you type  devpost run  in your terminal. "
    "Think of it as a behind-the-scenes tour: every file visited, every piece of data "
    "handed off, and every decision made — told as a story from first line of code to "
    "final Reddit post.",
    italic=True,
)
spacer()
doc.add_page_break()

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# CHAPTER 1 — THE CAST
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
h1("Chapter 1 — Meet the Characters")
body(
    "Before the story starts, here are the seven files that play a role. "
    "Each one has exactly one job."
)
spacer()

tbl = doc.add_table(rows=8, cols=2)
tbl.style = "Table Grid"
# header
hdr = tbl.rows[0].cells
hdr[0].text = "File"
hdr[1].text = "Its one job"
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = NAVY

rows_data = [
    ("main.py",         "The front door. Parses the  devpost run  command and passes control to the Agent."),
    ("config.py",       "The keychain. Reads API keys from  ~/.devpost/config.json  (or environment variables)."),
    ("post_log.py",     "The memory. Remembers which commit was last posted so we never repeat ourselves."),
    ("git_reader.py",   "The historian. Reads the Git log and turns raw commits into a clean summary string."),
    ("agent.py",        "The conductor. Orchestrates all the other files and decides what to do next."),
    ("tweet_builder.py","The tweet writer. Sends the Git summary to OpenAI and loops until the tweet fits in 280 chars."),
    ("reddit_poster.py","The community manager. Writes 6 different Reddit posts, one tailored for each subreddit."),
]
for i, (f, desc) in enumerate(rows_data, start=1):
    row = tbl.rows[i].cells
    row[0].paragraphs[0].add_run(f).font.name = "Courier New"
    row[1].text = desc

spacer()
doc.add_page_break()

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# CHAPTER 2 — THE STORY BEGINS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
h1("Chapter 2 — The Story Begins: You Type  devpost run")
body(
    "The moment you press Enter, the operating system finds the  devpost  binary "
    "and hands control to  main.py.  This file is nothing more than a traffic cop — "
    "it reads your flags (--path, --force, --dry-run) and then steps aside."
)
spacer()

h2("Scene 1 · main.py  ·  The Front Door")
code_block([
    "# What lives in main.py",
    "@cli.command()",
    "def run(path, force, dry_run):",
    "    config = ConfigManager()          # open the keychain",
    "    config.validate()                 # check keys exist",
    "    agent = DevPostAgent(config)      # create the conductor",
    "    agent.run(project_path=path, ...) # hand over control",
])
spacer()
body(
    "Two objects are created here and immediately passed deeper into the system:"
)
bullet("ConfigManager  — loads  ~/.devpost/config.json  into memory")
bullet("DevPostAgent   — receives the config so it can use the keys later")
spacer()
body(
    "After this moment,  main.py  does nothing else.  The whole story from here "
    "lives inside  agent.run().",
    italic=True,
    colour=TEAL,
)
spacer()

h2("Scene 2 · config.py  ·  Opening the Keychain")
body(
    "ConfigManager does one thing: read credentials.  "
    "It checks two places in order:"
)
bullet("First: environment variables (e.g.  OPENAI_API_KEY  — uppercase version of the key name)")
bullet("Second: the JSON file at  ~/.devpost/config.json")
body("The file looks like this:")
code_block([
    "{",
    '  "openai_api_key":       "sk-proj-...",',
    '  "reddit_client_id":     "abc123",',
    '  "reddit_client_secret": "xyz...",',
    '  "reddit_username":      "yourusername",',
    '  "reddit_password":      "yourpassword"',
    "}",
])
body(
    "If any of the five required keys are missing,  main.py  prints an error and stops "
    "right here — nothing else runs.  This is the only gate in the whole system."
)
spacer()
doc.add_page_break()

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# CHAPTER 3 — THE AGENT WAKES UP
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
h1("Chapter 3 — The Agent Wakes Up: agent.py Boots")
body(
    "DevPostAgent is the conductor of the orchestra.  When it is created, it wires "
    "up every specialist it might need:"
)
code_block([
    "class DevPostAgent:",
    "    def __init__(self, config):",
    "        self.client       = OpenAI(api_key=config.get('openai_api_key'))",
    "        self.post_log     = PostLog()          # the memory file",
    "        self.tweet_builder = TweetBuilder(client=self.client)",
    "        self.reddit_poster = RedditPoster(client=self.client, config=config)",
])
spacer()
body(
    "Notice that  TweetBuilder  and  RedditPoster  both receive the same  OpenAI client  "
    "object.  They never create their own — they borrow the one the Agent holds.  "
    "This is called dependency injection and it means the API key is set in exactly "
    "one place."
)
spacer()
body(
    "Now  agent.run()  begins its seven-step journey.",
    bold=True,
    colour=NAVY,
)
spacer()
doc.add_page_break()

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# CHAPTER 4 — THE SEVEN STEPS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
h1("Chapter 4 — Inside agent.run(): Seven Steps")

# STEP 1
h2("Step 1 of 7 — Read project_context.md")
body(
    "The Agent looks for a file called  project_context.md  in your project folder.  "
    "This file is written by you and describes your project in plain English: what it does, "
    "the tech stack, and the tone you want.  "
    "It is loaded as a raw string called  context."
)
code_block([
    "# agent.py — read_context()",
    "ctx_file = Path(project_path) / 'project_context.md'",
    "context  = ctx_file.read_text()   # → plain string, e.g. 1-2 KB of text",
])
body(
    "If the file does not exist,  context  is just an empty string  \"\"  and the AI "
    "has to guess from the Git commits alone.  It still works but posts will be more generic."
)
spacer()

# STEP 2
h2("Step 2 of 7 — Check the Memory, Then Ask Git")
body(
    "This is the most important decision in the whole system: "
    "which commits should we post about?"
)
body("The Agent asks  PostLog  one question:")
code_block([
    "last_hash = post_log.get_last_hash(project_path)",
    "# Returns either: 'a3f92bc...' (a full git commit hash)",
    "#              or: None (never posted from this repo before)",
])
spacer()
body("Two paths split from here:", bold=True)
bullet(
    "Path A — last_hash is None (first run, or --force flag used): "
    "ask GitReader for all commits in the last 24 hours."
)
bullet(
    "Path B — last_hash exists: ask GitReader for every commit that happened "
    "after that specific commit hash."
)
spacer()

h3("Inside git_reader.py — The Historian")
body(
    "GitReader opens your  .git  folder using a library called GitPython.  "
    "It walks the commit log from newest to oldest."
)
spacer()
body("Path A — time mode:", bold=True)
code_block([
    "def get_commits_last_hours(hours=24):",
    "    cutoff = now(UTC) - 24 hours",
    "    for commit in repo.iter_commits():   # newest first",
    "        if commit.time >= cutoff:",
    "            results.append(commit)",
    "        else:",
    "            break   # stop — everything older won't qualify",
])
spacer()
body("Path B — hash mode:", bold=True)
code_block([
    "def get_commits_since_hash(last_hash):",
    "    for commit in repo.iter_commits():   # newest first",
    "        if commit.hexsha == last_hash:",
    "            return results  # stop here — we already posted this one",
    "        results.append(commit)",
    "    # If we walked the entire log and never found the hash,",
    "    # it means a git rebase happened.  Fall back to 24-hour mode.",
])
spacer()
body("Each commit is turned into a small dictionary:")
code_block([
    "{",
    '  "hash":          "a3f92bc",',
    '  "full_hash":     "a3f92bc1d4e8...",',
    '  "message":       "add login page",',
    '  "author":        "Dhruv",',
    '  "timestamp":     "2026-05-14T10:30:00",',
    '  "files_changed": ["app.py", "templates/login.html"],',
    '  "insertions":    47,',
    '  "deletions":     3',
    "}",
])
spacer()
body(
    "The result is a  list[dict]  — a Python list where each item is one commit's dictionary.  "
    "This list travels back to the Agent."
)
spacer()

h3("Back in agent.py — Summarize")
body(
    "Once the commit list arrives,  GitReader.summarize_changes()  compresses it "
    "into a single readable string called  git_summary:"
)
code_block([
    "Total commits: 3",
    "",
    "Commits (newest first):",
    '  - a3f92bc: "add login page"',
    '  - b12cc01: "fix session bug"',
    '  - c44dd99: "update README"',
    "",
    "Files touched: app.py, README.md, templates/login.html",
    "Total changes: +112 lines added, -8 lines removed",
])
spacer()
body(
    "This plain-English string is the main ingredient that gets sent to the AI.  "
    "It also captures the newest hash ( a3f92bc... ) which will be saved at the end "
    "if something gets posted.",
    colour=GREEN,
    bold=True,
)
spacer()
doc.add_page_break()

# STEP 3
h2("Step 3 of 7 — Generate the Tweet")
body(
    "The Agent calls  tweet_builder.generate(git_summary, context).  "
    "Inside  tweet_builder.py,  a self-correction loop begins."
)
code_block([
    "# tweet_builder.py — the loop",
    "for attempt in range(1, 6):          # up to 5 tries",
    "    prompt = build_prompt(git_summary, context, previous_attempt, chars_over)",
    "    response = openai_client.chat.completions.create(",
    '        model="gpt-4o-mini",',
    "        max_tokens=300,",
    "        messages=[{'role': 'user', 'content': prompt}]",
    "    )",
    "    tweet = response.choices[0].message.content.strip()",
    "    if len(tweet) <= 280:",
    "        return tweet, len(tweet)      # ✅ valid — exit immediately",
    "    chars_over = len(tweet) - 280",
    "    previous_attempt = tweet          # send back for correction",
])
spacer()
body("What the first prompt says to the AI (simplified):")
bullet("You are writing for a student developer sharing daily progress.")
bullet("Tone: honest, learning in public, not corporate.")
bullet("Here is the project context: [project_context.md contents]")
bullet("Here is what was built from Git: [git_summary]")
bullet("Write ONE tweet under 280 characters with 2-3 hashtags.")
bullet("Only include a deployment URL for major milestones, not bug fixes.")
spacer()
body("What the retry prompt says:")
bullet("Your previous tweet was 23 characters too long (303 total).")
bullet("Previous attempt: [the tweet that failed]")
bullet("Rewrite it to be STRICTLY under 280 characters.")
bullet("Strategies: shorten phrases, cut one hashtag, simplify wording.")
spacer()
body(
    "The loop repeats until the tweet fits or 5 attempts are exhausted.  "
    "The user never sees a failed attempt — only the final valid tweet is returned.",
    italic=True,
    colour=TEAL,
)
body(
    "Return value back to agent.py:  (tweet_string, character_count)  — a tuple of two values."
)
spacer()
doc.add_page_break()

# STEP 4
h2("Step 4 of 7 — Generate 6 Reddit Posts")
body(
    "The Agent calls  reddit_poster.generate_all(git_summary, context).  "
    "This loops through six subreddits and makes one AI call per community."
)
code_block([
    "# reddit_poster.py",
    "SUBREDDIT_PERSONAS = {",
    '    "SideProject":      { "audience": "makers", "tone": "excited but genuine", ... },',
    '    "webdev":           { "audience": "web developers", "tone": "technical", ... },',
    '    "learnprogramming": { "audience": "beginners", "tone": "honest about struggles", ... },',
    '    "coding":           { "audience": "general programmers", "tone": "casual", ... },',
    '    "artificial":       { "audience": "AI enthusiasts", "tone": "technically curious", ... },',
    '    "MachineLearning":  { "audience": "ML researchers", "tone": "precise", ... },',
    "}",
])
spacer()
body("For each subreddit, the prompt says:")
bullet("Write a Reddit post for r/webdev.")
bullet("Audience: web developers interested in technical implementation.")
bullet("Tone: technical but approachable, share the stack and decisions.")
bullet("Here is the project context and Git activity: [...]")
bullet('Return ONLY a JSON object: {"title": "...", "body": "..."}')
spacer()
body(
    "The AI returns raw JSON.  reddit_poster.py parses it with  json.loads()  "
    "and stores the title and body for each subreddit.  "
    "The return value is a dictionary:"
)
code_block([
    "{",
    '  "SideProject":    { "title": "Built a login page today...", "body": "..." },',
    '  "webdev":         { "title": "Flask + Jinja2 session handling...", "body": "..." },',
    '  "learnprogramming": { "title": "Finally understood sessions...", "body": "..." },',
    '  ... (3 more)',
    "}",
])
body(
    "Six AI calls happen in sequence.  Each one is independent — "
    "if one fails, the others still run.  Any failure is recorded as  [generation failed]."
)
spacer()
doc.add_page_break()

# STEP 5
h2("Step 5 of 7 — Show Tweet, Ask for Approval")
body(
    "The Agent now calls  display.ask_tweet_approval(tweet, char_count).  "
    "This is the first moment a human sees anything."
)
body("The terminal shows a blue panel containing:")
bullet("The tweet text")
bullet("The character count coloured green (≤280) or red (>280)")
bullet("A yes/no question: 'Copy this tweet to clipboard?'")
spacer()
body("Two outcomes:")
bullet(
    "Yes → tweet_builder.copy_to_clipboard(tweet) is called.  "
    "pyperclip copies the tweet to your Mac clipboard.  "
    "You paste it yourself on X/Twitter.  "
    "self.anything_posted is set to True.",
    level=0
)
bullet(
    "No → tweet is skipped.  Results dictionary records  'skipped'.",
    level=0
)
spacer()

h2("Step 6 of 7 — Show Reddit Posts, Ask for Each")
body(
    "The Agent loops through all 6 subreddits.  For each one it calls  "
    "display.ask_reddit_approval()  which shows a magenta panel with the title and "
    "first 300 characters of the body."
)
body("Two outcomes per subreddit:")
bullet(
    "Yes → reddit_poster.post_to_subreddit(subreddit, title, body) is called.  "
    "PRAW (the Reddit library) sends the post to Reddit's API.  "
    "Reddit returns a URL like  reddit.com/r/webdev/comments/abc123/...  "
    "which is stored in results and printed to the terminal.",
    level=0
)
bullet(
    "No → subreddit skipped.  Results dictionary records  'skipped'.",
    level=0
)
spacer()
doc.add_page_break()

# STEP 7
h2("Step 7 of 7 — Save the Memory, Print Summary")
body(
    "After the approval loop finishes, the Agent checks one flag: "
    "was anything_posted set to True?"
)
code_block([
    "# agent.py — final step",
    "if self.anything_posted and newest_hash and not dry_run:",
    "    post_log.save_last_hash(project_path, newest_hash)",
    "    # Writes newest_hash into ~/.devpost/post_log.json",
    "    # keyed by the absolute path of your project folder",
])
spacer()
body("What gets written to  ~/.devpost/post_log.json:")
code_block([
    "{",
    '  "/Users/salescode/development/conversational-ai-assistant": "a3f92bc1d4e8..."',
    "}",
])
body(
    "Next time you run  devpost run,  PostLog reads this file and gives the hash to "
    "GitReader.  GitReader walks the log and stops the moment it sees that hash.  "
    "Only the commits that happened after it are included.  "
    "This is how the system never repeats itself.",
    colour=GREEN,
    bold=True,
)
spacer()
body("Finally,  display.print_final_summary(results)  prints a table:")
code_block([
    "┌────────────────────┬─────────────────────────────┐",
    "│ Action             │ Result                      │",
    "├────────────────────┼─────────────────────────────┤",
    "│ tweet              │ ✅ copied                    │",
    "│ r/SideProject      │ ✅ posted: reddit.com/...    │",
    "│ r/webdev           │ ⏭ skipped                   │",
    "│ r/learnprogramming │ ✅ posted: reddit.com/...    │",
    "│ r/coding           │ ⏭ skipped                   │",
    "│ r/artificial       │ ✅ posted: reddit.com/...    │",
    "│ r/MachineLearning  │ ❌ failed                    │",
    "└────────────────────┴─────────────────────────────┘",
])
spacer()
doc.add_page_break()

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# CHAPTER 5 — THE FULL TRAIL
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
h1("Chapter 5 — The Full Data Trail at a Glance")
body(
    "Here is every piece of data in the system, where it comes from, "
    "what shape it takes, and where it ends up."
)
spacer()

trail = [
    ("You type  devpost run",
     "main.py",
     "CLI flags: path, force, dry_run  (Python booleans/strings)"),

    ("main.py creates ConfigManager",
     "config.py  reads  ~/.devpost/config.json",
     "Python dict of credential strings"),

    ("main.py passes config into DevPostAgent",
     "agent.py  __init__",
     "OpenAI client object + PostLog + TweetBuilder + RedditPoster"),

    ("agent.py reads  project_context.md",
     "Filesystem  (your project folder)",
     "Plain string:  context  (1-2 KB describing your project)"),

    ("agent.py asks PostLog for last hash",
     "post_log.py  reads  ~/.devpost/post_log.json",
     "Either a full git hash string  or  None"),

    ("agent.py asks GitReader for commits",
     "git_reader.py  opens  .git/  via GitPython",
     "list[dict]  — one dict per commit (hash, message, files, +/- lines)"),

    ("agent.py calls summarize_changes",
     "git_reader.py  (static method)",
     "Plain string:  git_summary  (multi-line readable summary)"),

    ("agent.py calls tweet_builder.generate",
     "tweet_builder.py  →  OpenAI API  (gpt-4o-mini)",
     "Tuple: (tweet_string, int)  e.g.  ('Built login page today...', 247)"),

    ("agent.py calls reddit_poster.generate_all",
     "reddit_poster.py  →  OpenAI API  (gpt-4o-mini, 6 calls)",
     "dict[str, dict]  — subreddit name → {title, body}"),

    ("agent.py shows tweet, user approves",
     "display.py  →  terminal  →  pyperclip  (clipboard)",
     "Tweet string lands in your Mac clipboard  (no network call)"),

    ("agent.py shows each Reddit post, user approves",
     "reddit_poster.py  →  PRAW  →  Reddit API",
     "Reddit returns a URL string for each approved post"),

    ("agent.py saves post log",
     "post_log.py  writes  ~/.devpost/post_log.json",
     "newest_hash string stored under absolute project path key"),
]

tbl2 = doc.add_table(rows=len(trail)+1, cols=3)
tbl2.style = "Table Grid"
hdrs = tbl2.rows[0].cells
for c, text in zip(hdrs, ["Trigger", "File / Service", "Data Shape"]):
    c.paragraphs[0].add_run(text).bold = True
    for run in c.paragraphs[0].runs:
        run.font.color.rgb = NAVY

for i, (trigger, service, shape) in enumerate(trail, start=1):
    row = tbl2.rows[i].cells
    row[0].text = trigger
    row[1].text = service
    row[2].text = shape

spacer()
doc.add_page_break()

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# CHAPTER 6 — THE TWO FILES ON DISK
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
h1("Chapter 6 — The Two Persistent Files")
body(
    "Everything in Chapters 2-5 happens in RAM and disappears when the process exits.  "
    "These two files are the only things that survive between runs."
)
spacer()

h2("~/.devpost/config.json  — The Keychain")
body("Written once by  devpost setup.  Read by every run.  Never rewritten unless you re-run setup.")
code_block([
    "{",
    '  "openai_api_key":       "sk-proj-...",',
    '  "reddit_client_id":     "abc123",',
    '  "reddit_client_secret": "xyz...",',
    '  "reddit_username":      "yourusername",',
    '  "reddit_password":      "yourpassword"',
    "}",
])
spacer()

h2("~/.devpost/post_log.json  — The Memory")
body(
    "Written at the end of each run where something was posted.  "
    "Read at the start of every run to decide which commits are 'new'.  "
    "Keyed by the absolute path of your project folder so multiple projects are tracked independently."
)
code_block([
    "{",
    '  "/Users/salescode/development/conversational-ai-assistant": "a3f92bc1d4e8f7...",',
    '  "/Users/salescode/development/some-other-project":          "cc99de12ab56.."',
    "}",
])
spacer()
body(
    "If you delete this file, every project acts as if it's a first run and fetches "
    "the last 24 hours of commits.  This is what  devpost reset  does for the current folder.",
    italic=True,
    colour=ORANGE,
)
spacer()
doc.add_page_break()

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# CHAPTER 7 — WHAT HAPPENS WHEN THINGS GO WRONG
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
h1("Chapter 7 — What Happens When Things Go Wrong")

h2("No commits found")
body(
    "If GitReader returns an empty list, the Agent prints a friendly message "
    "('Nothing to post — write some code first!') and exits immediately.  "
    "No API calls are made.  The post log is unchanged."
)
spacer()

h2("Tweet is still too long after 5 retries")
body(
    "TweetBuilder falls back to a mechanical truncation: it takes the first 270 characters "
    "of git_summary, cuts at the last word boundary, and adds '...'.  "
    "This always fits but will be ugly — it rarely happens in practice."
)
spacer()

h2("OpenAI API error during generation")
body(
    "Each AI call is wrapped in try/except.  If it fails on attempt 1, the loop "
    "retries up to 5 times.  For Reddit posts, each subreddit is isolated — a failure "
    "on r/webdev doesn't affect r/MachineLearning."
)
spacer()

h2("Git hash not found (after a rebase)")
body(
    "If you rebased your repo and the saved hash no longer exists in the log, "
    "GitReader detects this (it walked the entire history without finding the hash) "
    "and automatically falls back to 24-hour mode.  You see a yellow warning in the terminal.  "
    "You may get some duplicate commits in this edge case."
)
spacer()

h2("Clipboard copy fails")
body(
    "On headless systems (no display server), pyperclip can't write to the clipboard.  "
    "The tweet is printed to the terminal with instructions to copy it manually.  "
    "self.anything_posted is still set to True so the post log is still saved."
)
spacer()

h2("Reddit not initialized")
body(
    "If  reddit_client_id  is missing from config (e.g. during  --dry-run),  "
    "the PRAW Reddit object is never created (self.reddit = None).  "
    "Calls to  post_to_subreddit  return None immediately with a warning.  "
    "No crash occurs."
)
spacer()
doc.add_page_break()

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# CHAPTER 8 — ONE-LINE SUMMARY OF EVERY FUNCTION
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
h1("Chapter 8 — Every Function, One Line Each")

fn_data = [
    # main.py
    ("main.py", "cli()",              "Click entry point — groups all devpost commands"),
    ("main.py", "run()",              "Validates credentials, creates Agent, calls agent.run()"),
    ("main.py", "setup()",            "Runs the interactive credential wizard"),
    ("main.py", "init()",             "Creates a project_context.md template in current folder"),
    ("main.py", "status()",           "Prints a table of all tracked repos and their last hash"),
    ("main.py", "reset()",            "Clears post log for current directory after confirmation"),
    # config.py
    ("config.py", "ConfigManager.__init__()", "Creates ~/.devpost dir and loads config.json into memory"),
    ("config.py", "_load()",                  "Reads config.json; returns {} on missing or corrupt file"),
    ("config.py", "_save()",                  "Writes current config dict back to config.json"),
    ("config.py", "get(key)",                 "Returns value from env var first, then config dict"),
    ("config.py", "set(key, value)",          "Saves one key-value pair and writes to disk immediately"),
    ("config.py", "validate()",               "Checks all 5 required keys exist; returns (bool, list)"),
    ("config.py", "setup_wizard()",           "Interactive Rich prompts for all credentials"),
    # post_log.py
    ("post_log.py", "PostLog.__init__()",        "Creates ~/.devpost dir and loads post_log.json"),
    ("post_log.py", "_normalize(repo_path)",     "Converts any path to absolute string (handles '.' etc.)"),
    ("post_log.py", "get_last_hash(repo_path)",  "Returns saved commit hash for this repo, or None"),
    ("post_log.py", "save_last_hash(path, hash)","Writes hash to log; trims to 40 entries if over 50"),
    ("post_log.py", "clear_log(repo_path)",      "Deletes this repo's entry from the log"),
    ("post_log.py", "get_all_entries()",         "Returns copy of entire log dict (used by devpost status)"),
    # git_reader.py
    ("git_reader.py", "GitReader.__init__()",             "Opens .git repo; raises ValueError if not a git repo"),
    ("git_reader.py", "_format_commit(commit)",           "Converts a GitPython commit object to a plain dict"),
    ("git_reader.py", "get_commits_since_hash(hash)",     "Returns all commits newer than the given hash"),
    ("git_reader.py", "get_commits_last_hours(hours)",    "Returns all commits from the last N hours"),
    ("git_reader.py", "get_newest_hash(commits)",         "Static: returns full_hash from the first commit in list"),
    ("git_reader.py", "summarize_changes(commits)",       "Static: formats commit list into a readable string"),
    # agent.py
    ("agent.py", "DevPostAgent.__init__()",  "Wires up OpenAI client, PostLog, TweetBuilder, RedditPoster"),
    ("agent.py", "read_context(path)",       "Reads project_context.md; warns if missing"),
    ("agent.py", "run(path, force, dry_run)","Orchestrates all 7 steps of the agent pipeline"),
    # tweet_builder.py
    ("tweet_builder.py", "TweetBuilder.__init__()", "Stores OpenAI client, sets model to gpt-4o-mini"),
    ("tweet_builder.py", "_count_chars(text)",      "Returns len(text) — counts every character"),
    ("tweet_builder.py", "_build_prompt(...)",       "Builds first-attempt or retry prompt based on whether previous_attempt is set"),
    ("tweet_builder.py", "generate(summary, ctx)",  "Runs the self-correction loop; returns (tweet, count)"),
    ("tweet_builder.py", "copy_to_clipboard(tweet)","Copies tweet via pyperclip; returns bool success"),
    # reddit_poster.py
    ("reddit_poster.py", "RedditPoster.__init__()",         "Inits OpenAI client and PRAW Reddit (deferred if no credentials)"),
    ("reddit_poster.py", "validate_credentials()",          "Calls reddit.user.me() to verify login works"),
    ("reddit_poster.py", "_generate_post(subreddit, ...)",  "One AI call → parses JSON → returns {title, body}"),
    ("reddit_poster.py", "generate_all(summary, context)",  "Calls _generate_post for all 6 subreddits"),
    ("reddit_poster.py", "post_to_subreddit(sub, title, body)", "Uses PRAW to submit the post; returns URL or None"),
    # display.py
    ("display.py", "print_header()",            "Shows the DevPost Agent banner"),
    ("display.py", "print_step(n, total, msg)", "Shows step progress like  [2/7] Checking post log..."),
    ("display.py", "print_thinking(msg)",       "Shows a dim italic robot message for AI activity"),
    ("display.py", "print_post_log_status()",   "Reports whether a previous hash was found"),
    ("display.py", "print_git_summary()",       "Shows git summary in a blue bordered panel"),
    ("display.py", "print_tweet_draft()",       "Shows tweet in a panel with character count"),
    ("display.py", "print_reddit_draft()",      "Shows Reddit post preview in a magenta panel"),
    ("display.py", "ask_tweet_approval()",      "Shows tweet panel + yes/no prompt; returns bool"),
    ("display.py", "ask_reddit_approval()",     "Shows Reddit panel + yes/no prompt; returns bool"),
    ("display.py", "print_post_log_saved()",    "Confirms the hash that was saved"),
    ("display.py", "print_success/error/warning()", "Coloured one-line status messages"),
    ("display.py", "print_final_summary()",     "Prints the results table at the end"),
]

tbl3 = doc.add_table(rows=len(fn_data)+1, cols=3)
tbl3.style = "Table Grid"
h_cells = tbl3.rows[0].cells
for c, t in zip(h_cells, ["File", "Function", "What it does"]):
    c.paragraphs[0].add_run(t).bold = True
    for run in c.paragraphs[0].runs:
        run.font.color.rgb = NAVY

for i, (f, fn, desc) in enumerate(fn_data, start=1):
    row = tbl3.rows[i].cells
    row[0].paragraphs[0].add_run(f).font.name = "Courier New"
    row[1].paragraphs[0].add_run(fn).font.name = "Courier New"
    row[2].text = desc

spacer()
doc.add_page_break()

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# CHAPTER 9 — THE END SCENE
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
h1("Chapter 9 — The End Scene")
body(
    "Here is the complete story in one paragraph:"
)
spacer()

para = doc.add_paragraph()
para.paragraph_format.space_after = Pt(8)
para.paragraph_format.left_indent  = Inches(0.3)
para.paragraph_format.right_indent = Inches(0.3)
run = para.add_run(
    "You type  devpost run.  main.py  wakes up, reads your API keys from  config.py, "
    "and hands control to  agent.py.  The Agent asks  post_log.py  whether it has seen this "
    "repo before — if yes, it gets the last commit hash; if no, it starts fresh.  "
    "It hands that information to  git_reader.py,  which opens your  .git  folder, "
    "walks backward through the commit history, and collects every commit that happened "
    "after the saved hash (or in the last 24 hours on a first run).  Those commits are "
    "compressed into a plain-English summary string.  The summary and your  project_context.md  "
    "are both passed to  tweet_builder.py,  which sends them to the OpenAI API and loops "
    "until the tweet fits in 280 characters.  The same summary and context go to  "
    "reddit_poster.py,  which makes six separate API calls — one per subreddit — each with "
    "a different persona prompt that matches the community's culture.  The Agent then pauses "
    "and shows you everything it created, asking for your approval on each item.  "
    "Approved tweets go to your clipboard via pyperclip.  Approved Reddit posts go live "
    "via the Reddit API through PRAW.  If anything was approved,  post_log.py  writes the "
    "newest commit hash to  ~/.devpost/post_log.json.  The next time you run the tool, "
    "it reads that hash and skips every commit you have already posted about."
)
run.font.size   = Pt(11)
run.font.color.rgb = GRAY
run.font.italic = True
spacer()

divider()
spacer()

final = doc.add_paragraph()
final.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = final.add_run("End of data flow story.")
r.font.size = Pt(12)
r.font.color.rgb = TEAL
r.font.italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = "/Users/salescode/development/post-agent/DevPost_DataFlow_Story.docx"
doc.save(out_path)
print(f"Saved → {out_path}")
